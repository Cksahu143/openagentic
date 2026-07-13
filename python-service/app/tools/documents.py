"""
Document and image tools. All operate on paths INSIDE a session's
workspace (see app/workspace/manager.py) — the agent writes/downloads a
file into its workspace first (e.g. via workspace_write_file or a future
download tool), then calls these to extract structured content from it.
"""
from __future__ import annotations

from pydantic import BaseModel, Field

from app.tools.registry import registry
from app.workspace.manager import WorkspaceError, get_workspace


class WorkspaceFileInput(BaseModel):
    user_id: str
    session_id: str
    path: str = Field(..., description="Path to the file, relative to the workspace root")


class WorkspaceFilePageInput(WorkspaceFileInput):
    max_pages: int = Field(20, ge=1, le=200)


@registry.register(
    "read_pdf",
    "Extract text (and basic table detection) from a PDF file in the workspace.",
    WorkspaceFilePageInput,
    requires_permission="workspace",
    timeout_s=30,
)
async def read_pdf(user_id: str, session_id: str, path: str, max_pages: int) -> dict:
    ws = get_workspace(user_id, session_id)
    try:
        abs_path = ws.resolve(path)
    except WorkspaceError as e:
        raise ValueError(str(e)) from e
    if not abs_path.is_file():
        raise ValueError(f"No such file: {path}")

    import pdfplumber

    pages_out = []
    with pdfplumber.open(str(abs_path)) as pdf:
        for i, page in enumerate(pdf.pages[:max_pages]):
            text = (page.extract_text() or "").strip()
            tables = page.extract_tables()
            pages_out.append(
                {"page": i + 1, "text": text[:4000], "table_count": len(tables)}
            )
    return {"path": path, "page_count_extracted": len(pages_out), "pages": pages_out}


@registry.register(
    "extract_pdf_images",
    "Extract embedded images from a PDF in the workspace into a subfolder, "
    "for cases where read_pdf's text extraction misses image-only content.",
    WorkspaceFilePageInput,
    requires_permission="workspace",
    timeout_s=30,
)
async def extract_pdf_images(user_id: str, session_id: str, path: str, max_pages: int) -> dict:
    ws = get_workspace(user_id, session_id)
    try:
        abs_path = ws.resolve(path)
    except WorkspaceError as e:
        raise ValueError(str(e)) from e
    if not abs_path.is_file():
        raise ValueError(f"No such file: {path}")

    import fitz  # PyMuPDF

    out_dir_name = f"{abs_path.stem}_images"
    saved: list[str] = []
    doc = fitz.open(str(abs_path))
    try:
        for page_index in range(min(len(doc), max_pages)):
            page = doc[page_index]
            for img_index, img in enumerate(page.get_images(full=True)):
                xref = img[0]
                base_image = doc.extract_image(xref)
                ext = base_image["ext"]
                rel_path = f"{out_dir_name}/page{page_index + 1}_img{img_index + 1}.{ext}"
                abs_out = ws.absolute_path(rel_path)
                with open(abs_out, "wb") as f:
                    f.write(base_image["image"])
                saved.append(rel_path)
    finally:
        doc.close()

    return {"path": path, "images_extracted": len(saved), "files": saved}


@registry.register(
    "read_docx",
    "Extract paragraph text and table contents from a .docx file in the workspace.",
    WorkspaceFileInput,
    requires_permission="workspace",
    timeout_s=20,
)
async def read_docx(user_id: str, session_id: str, path: str) -> dict:
    ws = get_workspace(user_id, session_id)
    try:
        abs_path = ws.resolve(path)
    except WorkspaceError as e:
        raise ValueError(str(e)) from e
    if not abs_path.is_file():
        raise ValueError(f"No such file: {path}")

    import docx  # python-docx

    document = docx.Document(str(abs_path))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    tables = [
        [[cell.text for cell in row.cells] for row in table.rows] for table in document.tables
    ]
    return {"path": path, "paragraphs": paragraphs[:500], "tables": tables[:20]}


@registry.register(
    "image_info",
    "Get basic metadata (dimensions, format) for an image file in the workspace. "
    "For semantic understanding of image CONTENT, use a multimodal model call, not this.",
    WorkspaceFileInput,
    requires_permission="workspace",
    timeout_s=10,
)
async def image_info(user_id: str, session_id: str, path: str) -> dict:
    ws = get_workspace(user_id, session_id)
    try:
        abs_path = ws.resolve(path)
    except WorkspaceError as e:
        raise ValueError(str(e)) from e
    if not abs_path.is_file():
        raise ValueError(f"No such file: {path}")

    from PIL import Image

    with Image.open(abs_path) as img:
        return {
            "path": path,
            "format": img.format,
            "width": img.width,
            "height": img.height,
            "mode": img.mode,
        }
